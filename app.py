import os
import streamlit as st
import re
from docx import Document
from fpdf import FPDF
import PyPDF2
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env
load_dotenv()

# Configure Gemini API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Predefined list of skills for regex matching
PREDEFINED_SKILLS = [
    "Python", "Java", "SQL", "Machine Learning", "Data Analysis",
    "Data Engineering", "AWS", "Azure", "Docker", "Kubernetes",
    "ETL", "Big Data", "Hadoop", "Spark", "Tableau", "Power BI",
    "AWS Glue", "PySpark", "Aurora DB", "Dynamo DB", "Redshift",
    "Data Warehousing", "CI/CD", "Stone branch", "Scheduling Tool"
]

# Function to extract multi-word skills using regex
def match_skill(text, skills_list):
    matched_skills = set()
    for skill in skills_list:
        if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE):
            matched_skills.add(skill)
    return matched_skills

# Function to extract skills using Gemini API
def gemini_extract_skills(text):
    try:
        # Configure the model
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        # Initialize the model
        model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config,
        )

        # Start a chat session
        chat_session = model.start_chat(history=[])

        # Create the prompt
        prompt = (
            "Extract all skills from the following job description text. "
            "Ensure the response contains only skill names, separated by commas:\n\n"
            f"{text}"
        )

        # Send the prompt to Gemini
        response = chat_session.send_message(prompt)

        # Parse and clean the response
        skills_text = response.text.strip()
        skills = {skill.strip() for skill in skills_text.split(",") if skill.strip()}
        return skills

    except Exception as e:
        st.error(f"Failed to extract skills using Gemini API: {e}")
        return set()

# Function to extract text from a Word document
def extract_text_from_docx(docx_file):
    text = ''
    try:
        doc = Document(docx_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
    return text

# Function to extract text from a PDF document
def extract_text_from_pdf(pdf_file):
    text = ''
    try:
        reader = PyPDF2.PdfReader(pdf_file)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
    return text

# Function to combine skills from regex and Gemini
def extract_skills(job_description):
    # Extract using regex
    regex_skills = match_skill(job_description, PREDEFINED_SKILLS)

    # Extract using Gemini
    gemini_skills = gemini_extract_skills(job_description)

    # Combine results
    all_skills = regex_skills.union(gemini_skills)
    return list(all_skills)

# Function to create a Word document
def create_word_file(skills, job_name):
    doc = Document()
    doc.add_heading('Extracted Skills', level=1)
    for skill in skills:
        doc.add_paragraph(skill)
    file_path = f"{job_name}_Skills.docx"
    doc.save(file_path)
    return file_path

# Function to create a PDF file
def create_pdf_file(skills, job_name):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Extracted Skills", ln=True, align='C')
    for skill in skills:
        pdf.cell(200, 10, txt=skill, ln=True)
    file_path = f"{job_name}_Skills.pdf"
    pdf.output(file_path)
    return file_path

# Utility function to sanitize filenames
def sanitize_filename(filename):
    return "".join(c if c.isalnum() or c in " ._-()" else "_" for c in filename)

# Streamlit app
def main():
    st.title("Job Description Skills Parser")

    # File uploader for job description
    uploaded_file = st.file_uploader("Upload Job Description (Word or PDF)", type=["docx", "pdf"])

    if st.button("Extract Skills"):
        if uploaded_file:
            # Extract text based on file type
            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                job_description = extract_text_from_docx(uploaded_file)
                job_name = os.path.splitext(uploaded_file.name)[0]
            elif uploaded_file.type == "application/pdf":
                job_description = extract_text_from_pdf(uploaded_file)
                job_name = os.path.splitext(uploaded_file.name)[0]
            else:
                st.error("Unsupported file type.")
                return

            if job_description.strip():  # Ensure extracted text is not empty
                skills = extract_skills(job_description)
                if skills:
                    st.success("Skills extracted successfully!")
                    st.write("### Extracted Skills:")
                    st.write(", ".join(skills))  # Display the extracted skills

                    # Generate downloadable Word file
                    word_file_path = create_word_file(skills, job_name)
                    with open(word_file_path, "rb") as f:
                        st.download_button(
                            label="Download Skills as Word",
                            data=f,
                            file_name=f"{sanitize_filename(job_name)}_Skills.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    os.remove(word_file_path)  # Clean up the generated file

                    # Generate downloadable PDF file
                    pdf_file_path = create_pdf_file(skills, job_name)
                    with open(pdf_file_path, "rb") as f:
                        st.download_button(
                            label="Download Skills as PDF",
                            data=f,
                            file_name=f"{sanitize_filename(job_name)}_Skills.pdf",
                            mime="application/pdf"
                        )
                    os.remove(pdf_file_path)  # Clean up the generated file
                else:
                    st.warning("No skills found in the job description.")
            else:
                st.error("Failed to extract text from the uploaded document.")
        else:
            st.error("Please upload a job description.")

if __name__ == "__main__":
    main()
